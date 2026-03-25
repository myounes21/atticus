from pathlib import Path

import pytest
from docx import Document

from backend.ingestion.parsers.docx_parser import DocxParser


def test_docx_parser_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"

    document = Document()
    document.add_paragraph("Contract summary")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Clause"
    table.rows[0].cells[1].text = "Status"
    table.rows[1].cells[0].text = "Termination"
    table.rows[1].cells[1].text = "Pending"
    document.save(str(file_path))

    parsed = DocxParser().parse(file_path)

    assert parsed.metadata.document_name == "sample.docx"
    assert parsed.metadata.file_type == "docx"
    assert "Contract summary" in parsed.text
    assert "Clause | Status" in parsed.text


def test_docx_parser_raises_for_empty_document(tmp_path: Path) -> None:
    file_path = tmp_path / "empty.docx"

    document = Document()
    document.save(str(file_path))

    with pytest.raises(ValueError, match="DOCX contains no extractable text"):
        DocxParser().parse(file_path)


